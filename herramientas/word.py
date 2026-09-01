#!/usr/bin/env python3
"""Vuelca el volumen —o un trozo— a .docx con estilos de Word con nombre.

El PDF lo maqueta `libro.py` con su CSS. Este otro camino existe para poder
**dar formato a mano en Word** y devolver el resultado: cada cosa del volumen
lleva aquí un **estilo con nombre**, así que cambiar el estilo en Word cambia
todas sus apariciones de golpe, y el fichero devuelto dice, en `word/styles.xml`,
qué tipografía y qué medidas se han elegido para cada uno. Eso es lo que se
traslada después al CSS de `libro.py`.

Los estilos son deliberadamente pocos. Los de encabezado y lista son los
**integrados** de Word —«Heading 1», «List Bullet»…—, que cada Word muestra en
su idioma y que ya conocen el panel de estilos y el índice automático. Los
propios del temario van en castellano y son estos:

    Rótulo               el renglón en versalita de encima de cada título
    Portada dato         los tres renglones del pie de la portada
    Aviso                los párrafos de la caja «Cómo usar este volumen»
    Índice tema          el renglón de cada tema en el índice general
    Índice epígrafe      el de cada epígrafe debajo
    Cita de la norma     el texto literal del BOE (los blockquote del tema)
    Tabla                el texto dentro de las celdas
    Tabla cabecera       el de la primera fila
    Pregunta             el enunciado de una pregunta real de examen
    Opción               cada una de las cuatro respuestas
    Fuente de pregunta   el cuadernillo y el número, debajo de cada pregunta

**No hay formato a mano en ninguna parte**: si algo se ve distinto es porque
tiene un estilo distinto. Eso es lo que permite leer el fichero devuelto y
trasladar cada decisión al CSS de `libro.py` sin adivinar.

Uso:  word.py [salida.docx] [--temas 1,2,3] [--muestrario]
      python3 herramientas/word.py libro-tema-1.docx --temas 1 --muestrario

`--muestrario` añade al final una hoja con **todos los estilos, uno debajo de
otro y con su nombre delante**. No es parte del volumen: está para dar el
formato de una sentada y ver de un vistazo qué queda sin decidir.
"""
import os
import re
import sys
from datetime import date

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from markdown_it import MarkdownIt

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from libro import (CORTE, ERRATAS, RAIZ, TEMAS, lee, numera, ordena_opciones,
                   preguntas, sin_marcas)

md = MarkdownIt("commonmark").enable("table").enable("strikethrough")

GRIS = RGBColor(0x66, 0x66, 0x66)
TINTA = RGBColor(0x5A, 0x6B, 0x86)
ENCABEZADO = "TOAC – Temarios de Oposiciones"
PIE_DERECHA = "Oposiciones RTVE – Temario General"
NEGRO = RGBColor(0x11, 0x11, 0x11)


# ── estilos ──────────────────────────────────────────────────────────────────

def campo(parrafo, instruccion):
    """Mete un campo de Word en el párrafo: lo calcula Word, no nosotros.

    Es lo que hace que el índice traiga el número de página **de verdad** y que
    se pueda pinchar, y que el pie diga «Página 3 de 254» sin que nadie cuente.
    """
    from docx.oxml import OxmlElement
    r = parrafo.add_run()._r
    ini = OxmlElement("w:fldChar"); ini.set(qn("w:fldCharType"), "begin")
    txt = OxmlElement("w:instrText"); txt.set(qn("xml:space"), "preserve")
    txt.text = instruccion
    sep = OxmlElement("w:fldChar"); sep.set(qn("w:fldCharType"), "separate")
    hueco = OxmlElement("w:t"); hueco.set(qn("xml:space"), "preserve"); hueco.text = " "
    fin = OxmlElement("w:fldChar"); fin.set(qn("w:fldCharType"), "end")
    for e in (ini, txt, sep, hueco, fin):
        r.append(e)
    return parrafo


def actualiza_al_abrir(doc):
    """Pide a Word que recalcule los campos al abrir el documento.

    Sin esto el índice aparece vacío hasta que alguien pulsa F9, y quien reciba
    el fichero no tiene por qué saberlo.
    """
    from docx.oxml import OxmlElement
    e = OxmlElement("w:updateFields")
    e.set(qn("w:val"), "true")
    # el orden de los hijos de `w:settings` lo fija el esquema: `updateFields` va
    # delante de `compat`, no al final
    doc.settings.element.insert_element_before(
        e, "w:hdrShapeDefaults", "w:footnotePr", "w:endnotePr", "w:compat",
        "w:docVars", "w:rsids", "m:mathPr", "w:attachedSchema", "w:themeFontLang",
        "w:clrSchemeMapping", "w:doNotIncludeSubdocsInStats",
        "w:doNotAutoCompressPictures", "w:forceUpgrade", "w:captions",
        "w:readModeInkLockDown", "w:smartTagType", "w:shapeDefaults",
        "w:doNotEmbedSmartTags", "w:decimalSymbol", "w:listSeparator")


def estilo_parrafo(doc, nombre, base="Normal", **kw):
    """Crea un estilo de párrafo si no existe y le pone unos valores de salida.

    Los valores son un punto de partida, no una decisión: están para que el
    documento se lea antes de tocarlo. Lo que valga de verdad lo dirá el
    fichero que vuelva de Word.
    """
    if nombre in [s.name for s in doc.styles]:
        return doc.styles[nombre]
    st = doc.styles.add_style(nombre, 1)  # 1 = WD_STYLE_TYPE.PARAGRAPH
    st.base_style = doc.styles[base]
    f, p = st.font, st.paragraph_format
    if "tam" in kw:
        f.size = Pt(kw["tam"])
    if kw.get("cursiva"):
        f.italic = True
    if kw.get("negrita"):
        f.bold = True
    if kw.get("versalita"):
        f.all_caps = True
        _espacia_letras(f.element.get_or_add_rPr(), 24)
    if "color" in kw:
        f.color.rgb = kw["color"]
    if "tipo" in kw:
        f.name = kw["tipo"]
    p.space_before = Pt(kw.get("antes", 0))
    p.space_after = Pt(kw.get("despues", 4))
    if kw.get("centrado"):
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if kw.get("izquierda"):
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if "sangria" in kw:
        p.left_indent = Cm(kw["sangria"])
    if kw.get("juntar"):
        p.keep_with_next = True
    return st


def _arregla_zoom(doc):
    """La plantilla de python-docx trae `<w:zoom>` sin el atributo `w:percent`,
    que el esquema exige. No se ve, pero deja el fichero inválido."""
    z = doc.settings.element.find(qn("w:zoom"))
    if z is not None and z.get(qn("w:percent")) is None:
        z.set(qn("w:percent"), "100")


def _espacia_letras(rPr, veinteavos):
    """Separa las letras del rótulo. El orden de los hijos de `w:rPr` lo fija el
    esquema, así que `w:spacing` no se puede añadir al final: va detrás de
    `w:color` y delante de `w:sz`."""
    from docx.oxml import OxmlElement
    e = OxmlElement("w:spacing")
    e.set(qn("w:val"), str(veinteavos))
    rPr.insert_element_before(e, "w:w", "w:kern", "w:position", "w:sz", "w:szCs",
                              "w:highlight", "w:u", "w:effect", "w:bdr", "w:shd",
                              "w:fitText", "w:vertAlign", "w:rtl", "w:cs", "w:em",
                              "w:lang", "w:eastAsianLayout", "w:specVanish", "w:oMath")


def cabecera_y_pie(doc):
    """Encabezado y pie en todas las páginas menos la portada.

    Word lo resuelve con «primera página distinta»: la portada es la primera de
    la sección, así que se queda limpia dejando vacíos su encabezado y su pie.
    """
    sec = doc.sections[0]
    sec.different_first_page_header_footer = True
    sec.header_distance = Cm(1.2)
    sec.footer_distance = Cm(1.2)

    enc = sec.header.paragraphs[0]
    enc.style = doc.styles["Encabezado TOAC"]
    enc.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    enc.add_run(ENCABEZADO)

    pie = sec.footer.paragraphs[0]
    pie.style = doc.styles["Pie"]
    pie.paragraph_format.tab_stops.add_tab_stop(
        Cm(16.6), WD_TAB_ALIGNMENT.RIGHT)
    pie.add_run("Página ")
    campo(pie, " PAGE ")
    pie.add_run(" de ")
    campo(pie, " NUMPAGES ")
    pie.add_run("\t" + PIE_DERECHA)


def prepara(doc):
    """Deja el documento con la página, el cuerpo y los estilos propios."""
    s = doc.sections[0]
    s.page_width, s.page_height = Cm(21), Cm(29.7)
    s.top_margin = s.bottom_margin = Cm(2)
    s.left_margin = s.right_margin = Cm(2.2)

    n = doc.styles["Normal"]
    n.font.name = "Georgia"
    n.font.size = Pt(10.5)
    n.font.color.rgb = NEGRO
    n.paragraph_format.space_after = Pt(6)
    n.paragraph_format.line_spacing = 1.3
    n.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for i, tam in ((1, 20), (2, 14), (3, 11.5), (4, 10.5)):
        h = doc.styles["Heading %d" % i]
        h.font.name = "Helvetica Neue"
        h.font.size = Pt(tam)
        h.font.color.rgb = NEGRO
        h.font.bold = i <= 2
        h.font.italic = i == 4
        h.paragraph_format.keep_with_next = True
        h.paragraph_format.space_before = Pt(14 if i <= 2 else 10)
        h.paragraph_format.space_after = Pt(4)
    # Title y Subtitle solo salen en la portada, así que van centrados en el
    # propio estilo y no a mano sobre el párrafo
    for nombre, tipo, tam in (("Title", "Helvetica Neue", 30), ("Subtitle", "Georgia", 13)):
        st = doc.styles[nombre]
        st.font.name, st.font.size = tipo, Pt(tam)
        st.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _arregla_zoom(doc)
    actualiza_al_abrir(doc)

    estilo_parrafo(doc, "Rótulo", tam=8.5, color=GRIS, versalita=True,
                   tipo="Helvetica Neue", juntar=True, despues=2)
    estilo_parrafo(doc, "Ficha", tam=8.5, tipo="Helvetica Neue", despues=1)
    estilo_parrafo(doc, "Cita de la norma", tam=9.5, sangria=0.8, antes=4, despues=4)
    estilo_parrafo(doc, "Pregunta", tam=9.5, despues=2, juntar=True)
    estilo_parrafo(doc, "Opción", tam=9.5, sangria=0.7, despues=1)
    estilo_parrafo(doc, "Fuente de pregunta", tam=7.5, color=GRIS,
                   tipo="Helvetica Neue", sangria=0.7, despues=8)
    estilo_parrafo(doc, "Aviso", tam=10, despues=6)
    estilo_parrafo(doc, "Portada dato", tam=10, despues=10, centrado=True)
    estilo_parrafo(doc, "Portada rótulo", base="Rótulo", centrado=True, despues=6)
    estilo_parrafo(doc, "Separador", despues=8, antes=8)
    estilo_parrafo(doc, "Índice tema", tam=10.5, negrita=True, antes=6, despues=2,
                   juntar=True)
    estilo_parrafo(doc, "Índice epígrafe", tam=9.5, color=GRIS, sangria=0.8, despues=0)
    estilo_parrafo(doc, "Tabla", tam=9, despues=1, izquierda=True)
    estilo_parrafo(doc, "Tabla cabecera", base="Tabla", negrita=True)
    estilo_parrafo(doc, "Encabezado TOAC", tam=9, color=TINTA, despues=0)
    estilo_parrafo(doc, "Pie", tam=8.5, color=TINTA, despues=0)
    for nombre, ajusta in (("Número de pregunta", lambda f: (setattr(f, "bold", True),
                                                             setattr(f, "name", "Helvetica Neue"))),
                           ("Código", lambda f: (setattr(f, "name", "Consolas"),
                                                 setattr(f, "size", Pt(9))))):
        if nombre not in [x.name for x in doc.styles]:
            ajusta(doc.styles.add_style(nombre, 2).font)  # 2 = CHARACTER
    return doc


# ── del markdown al documento ────────────────────────────────────────────────

def escribe_inline(p, tok):
    """Vuelca los hijos de un token inline como runs con su formato."""
    negrita = cursiva = codigo = 0
    for h in tok.children or []:
        if h.type == "strong_open":
            negrita += 1
        elif h.type == "strong_close":
            negrita -= 1
        elif h.type == "em_open":
            cursiva += 1
        elif h.type == "em_close":
            cursiva -= 1
        elif h.type == "softbreak":
            p.add_run(" ")
        elif h.type in ("text", "code_inline"):
            codigo = h.type == "code_inline"
            r = p.add_run(h.content, style="Código" if codigo else None)
            r.bold = negrita > 0
            r.italic = cursiva > 0
        # los enlaces se dejan como texto: en papel una URL azul subrayada
        # no lleva a ninguna parte, y la trazabilidad ya está en la ficha


GRIS_TABLA = "E4E4E4"


def sombrea(celda, color):
    """Pinta el fondo de una celda."""
    from docx.oxml import OxmlElement
    e = OxmlElement("w:shd")
    e.set(qn("w:val"), "clear")
    e.set(qn("w:color"), "auto")
    e.set(qn("w:fill"), color)
    celda._tc.get_or_add_tcPr().append(e)


def escribe_tabla(doc, filas, cabecera, alterna=False):
    """Escribe una tabla.

    **La primera fila va en gris siempre**, lleve cabecera o no. Y con
    `alterna`, el gris se repite en las filas pares: es lo que hace legible el
    solucionario, donde una fila son los números y la siguiente sus respuestas.
    """
    t = doc.add_table(rows=0, cols=max(len(f) for f in filas))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, fila in enumerate(filas):
        celdas = t.add_row().cells
        gris = i == 0 or (alterna and i % 2 == 0)
        for j, tok in enumerate(fila):
            if j >= len(celdas):
                continue
            celdas[j].text = ""
            p = celdas[j].paragraphs[0]
            p.style = doc.styles["Tabla cabecera" if i == 0 and cabecera else "Tabla"]
            escribe_inline(p, tok)
            if gris:
                sombrea(celdas[j], GRIS_TABLA)
    return t


ANCLA = re.compile(r'<a id="[^"]+"></a>\s*')


def vuelca(doc, texto, salto_titulos=0):
    texto = ANCLA.sub("", texto)
    toks = md.parse(texto)
    i, pilas = 0, []          # pilas: ("bullet"|"ordered", nivel)
    while i < len(toks):
        t = toks[i]
        if t.type == "heading_open":
            nivel = min(int(t.tag[1]) + salto_titulos, 4)
            p = doc.add_paragraph(style="Heading %d" % nivel)
            escribe_inline(p, toks[i + 1])
            i += 3
            continue
        if t.type == "paragraph_open":
            estilo = "Normal"
            if pilas:
                clase, nivel = pilas[-1]
                estilo = ("List Bullet" if clase == "bullet" else "List Number")
                if nivel > 1:
                    estilo += " %d" % min(nivel, 3)
            elif t.level and toks[max(i - 1, 0)].type != "blockquote_open":
                pass
            p = doc.add_paragraph(style=estilo)
            escribe_inline(p, toks[i + 1])
            i += 3
            continue
        if t.type == "blockquote_open":
            j, dentro = i + 1, []
            hondo = 1
            while j < len(toks) and hondo:
                if toks[j].type == "blockquote_open":
                    hondo += 1
                elif toks[j].type == "blockquote_close":
                    hondo -= 1
                    if not hondo:
                        break
                dentro.append(toks[j])
                j += 1
            for k, x in enumerate(dentro):
                if x.type == "inline":
                    p = doc.add_paragraph(style="Cita de la norma")
                    escribe_inline(p, x)
            i = j + 1
            continue
        if t.type in ("bullet_list_open", "ordered_list_open"):
            clase = "bullet" if t.type.startswith("bullet") else "ordered"
            pilas.append((clase, len(pilas) + 1))
            i += 1
            continue
        if t.type in ("bullet_list_close", "ordered_list_close"):
            if pilas:
                pilas.pop()
            i += 1
            continue
        if t.type == "table_open":
            j, filas, fila, cabecera = i + 1, [], [], False
            while j < len(toks) and toks[j].type != "table_close":
                if toks[j].type == "thead_open":
                    cabecera = True
                elif toks[j].type in ("th_open", "td_open"):
                    fila.append(toks[j + 1])
                elif toks[j].type == "tr_close":
                    filas.append(fila)
                    fila = []
                j += 1
            if filas:
                escribe_tabla(doc, filas, cabecera)
                doc.add_paragraph(style="Normal")
            i = j + 1
            continue
        if t.type == "hr":
            _raya(doc.add_paragraph(style="Separador"))
            i += 1
            continue
        i += 1


def _raya(p):
    from docx.oxml import OxmlElement
    pPr = p._p.get_or_add_pPr()
    bordes = OxmlElement("w:pBdr")
    b = OxmlElement("w:bottom")
    b.set(qn("w:val"), "single")
    b.set(qn("w:sz"), "4")
    b.set(qn("w:color"), "BBBBBB")
    bordes.append(b)
    pPr.append(bordes)


def salto(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


# ── el volumen ───────────────────────────────────────────────────────────────

def parte_tema(doc, numero, base, banco):
    crudo = sin_marcas(lee("temas/general/%s.md" % base))
    titulo = re.search(r"(?m)^# (.+)$", crudo).group(1)
    cuerpo = re.sub(r"(?m)^# .+$\n", "", crudo, count=1)
    cuerpo = re.sub(r"## Índice\n.*?(?=\n## )", "", cuerpo, flags=re.S)

    ficha, resto = None, cuerpo
    mt = re.match(r"\s*(\|.*?\|)\n\n", cuerpo, re.S)
    if mt:
        ficha, resto = mt.group(1), cuerpo[mt.end():]

    resto, _ = numera(resto, numero)
    salto(doc)
    doc.add_paragraph("Temario general", style="Rótulo")
    doc.add_paragraph("TEMA %d – %s" % (numero, titulo.split("·", 1)[-1].strip()),
                      style="Heading 1")
    if ficha:
        filas = []
        for linea in ficha.strip().split("\n"):
            if set(linea.replace("|", "").strip()) <= set("- :"):
                continue
            celdas = [c.strip() for c in linea.strip().strip("|").split("|")]
            filas.append([md.parse(c)[1] for c in celdas])
        escribe_tabla(doc, filas, False)
        doc.add_paragraph(style="Normal")
    vuelca(doc, resto)

    esquema = sin_marcas(lee("esquemas/general/%s.md" % base))
    esquema = re.sub(r"(?m)^# .+$\n", "", esquema, count=1)
    esquema = re.sub(r"## Índice\n.*?(?=\n## )", "", esquema, flags=re.S)
    salto(doc)
    doc.add_paragraph("Para repasar", style="Rótulo")
    doc.add_paragraph("Esquema de repaso · tema %d" % numero, style="Heading 1")
    vuelca(doc, esquema)

    if not banco:
        return []
    ps = preguntas(banco)
    salto(doc)
    doc.add_paragraph("Para comprobar si el tema se sostiene", style="Rótulo")
    doc.add_paragraph("Preguntas reales de examen · %s"
                      % ("temas 2 y 3" if banco == "g2-g3" else "tema %d" % numero),
                      style="Heading 1")
    p = doc.add_paragraph(style="Normal")
    p.add_run("%d preguntas de los cuadernillos de 2024. Las respuestas, al final "
              "del volumen." % len(ps)).italic = True
    for n, (_, enunciado, _) in enumerate(ps, 1):
        pinta_pregunta(doc, n, enunciado)
    return ps


def pinta_pregunta(doc, n, enunciado):
    cabeza, opciones, sin_texto = ordena_opciones(enunciado)
    p = doc.add_paragraph(style="Pregunta")
    p.add_run("%d. " % n, style="Número de pregunta")
    p.add_run(cabeza).bold = True
    for o in opciones:
        doc.add_paragraph(o, style="Opción")
    if sin_texto:
        doc.add_paragraph(
            "El examen corta estas opciones sin punto final, así que la transcripción "
            "no marca dónde acaba cada una y %s se queda sin texto. Se imprime como salió."
            % ", ".join("la %s)" % l for l in sin_texto), style="Fuente de pregunta")


def portada(doc, total_preg, cuantos):
    doc.add_paragraph("Oposiciones RTVE · convocatorias 1/2022 y 3/2022",
                      style="Portada rótulo")
    doc.add_paragraph("Temario general", style="Title")
    doc.add_paragraph("Los ocho temas comunes a Producción (Asistencia), "
                      "Documentación e Información y Contenidos", style="Subtitle")
    cuerpo = ("Ocho temas · ocho esquemas de repaso" if cuantos == 8 else
              "Tema %d de ocho, con su esquema de repaso"
              % cuantos if cuantos == 1 else "%d de los ocho temas" % cuantos)
    for l in ("Redacción vigente a %s" % CORTE,
              "%s · %d preguntas reales de examen" % (cuerpo, total_preg),
              "Generado el %s" % date.today().strftime("%d/%m/%Y")):
        doc.add_paragraph(l, style="Portada dato")


def aviso(doc):
    salto(doc)
    doc.add_paragraph("Cómo usar este volumen", style="Heading 1")
    caja = doc.add_paragraph(style="Aviso")
    caja.add_run("La redacción que vale es la del %s" % CORTE).bold = True
    caja.add_run(", que es la fecha de corte que imponen las bases: «las pruebas se "
                 "realizarán sobre su texto vigente a fecha de la primera publicación de "
                 "las Bases Generales». Lo que cambió después está en el tema, en "
                 "apartados marcados como ")
    caja.add_run("notas de actualización").italic = True
    caja.add_run(", y no es materia examinable.")
    _recuadra(caja)
    for trozos in (
        [("Cada tema trae tres partes.", True),
         (" El cuerpo, para leer; el esquema, para repasar, que va detrás y no delante a "
          "propósito; y las preguntas reales de los cuadernillos de 2024, para comprobar "
          "si el tema se sostiene. Las respuestas están al final del volumen, no junto a "
          "la pregunta: con la respuesta a la vista no hay autoevaluación.", False)],
        [("Tres respuestas oficiales de 2024 están mal.", True),
         (" Van marcadas una a una en el apéndice, con el precepto que las desmiente. El "
          "temario enseña la norma, no la plantilla.", False)],
        [("Las preguntas se imprimen tal como salieron del cuadernillo", True),
         (", sin más limpieza que quitarles el pie de página. Están leídas una a una: las "
          "que la clasificación por palabras clave había puesto en el tema que no les "
          "tocaba —o que no son del temario general— se recolocaron contra la fuente.",
          False)],
        [("Nada de aquí se ha escrito de memoria.", True),
         (" Cada dato se ha leído en el texto consolidado del BOE en su redacción a la "
          "fecha de corte, o en la fuente oficial que se cita en la trazabilidad de cada "
          "tema.", False)]):
        p = doc.add_paragraph(style="Aviso")
        for texto, negrita in trozos:
            p.add_run(texto).bold = negrita


def _recuadra(p):
    from docx.oxml import OxmlElement
    pPr = p._p.get_or_add_pPr()
    bordes = OxmlElement("w:pBdr")
    for lado in ("top", "left", "bottom", "right"):
        b = OxmlElement("w:%s" % lado)
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "8")
        b.set(qn("w:space"), "8")
        b.set(qn("w:color"), "111111")
        bordes.append(b)
    pPr.append(bordes)


def indice(doc):
    """El índice, como campo de Word: con página, clicable y siempre al día.

    No se escribe a mano. Se le pide a Word que lo arme con los encabezados de
    los niveles 1 a 4, que son los que numera `numera()`: así trae el número de
    página de verdad, se puede pinchar para saltar, y **se rehace solo** si el
    documento crece al cambiarle el formato.
    """
    salto(doc)
    doc.add_paragraph("Índice general", style="Heading 1")
    p = doc.add_paragraph(style="Normal")
    p.add_run("Si el índice aparece vacío, se rellena con ")
    p.add_run("Ctrl+E y luego F9").bold = True
    p.add_run(" —o pulsando «Sí» cuando Word pregunte al abrir si actualiza los "
              "campos—. Cada entrada lleva su página y se puede pinchar para ir a ella.")
    campo(doc.add_paragraph(style="Normal"), r' TOC \o "1-4" \h \z \u ')


def respuestas(doc, hechos):
    salto(doc)
    doc.add_paragraph("Apéndice", style="Rótulo")
    doc.add_paragraph("Respuestas oficiales", style="Heading 1")
    p = doc.add_paragraph(style="Normal")
    p.add_run("La respuesta es la de la ")
    p.add_run("plantilla oficial").bold = True
    p.add_run(" del examen, y el número es el que la pregunta lleva impreso en su tema. ")
    p.add_run("Tres respuestas oficiales están mal").bold = True
    p.add_run(" y van avisadas debajo de su tabla: el volumen enseña la norma, no la "
              "plantilla.")
    for numero, banco, ps in hechos:
        doc.add_paragraph("Temas 2 y 3" if banco == "g2-g3" else "Tema %d" % numero,
                          style="Heading 2")
        sueltas = [(n, r) for n, (_, _, r) in enumerate(ps, 1)]
        POR_FILA = 10
        filas = []
        for a in range(0, len(sueltas), POR_FILA):
            trozo = sueltas[a:a + POR_FILA]
            filas.append([md.parse("**%d**" % n)[1] for n, _ in trozo])
            filas.append([md.parse("**%s**" % r)[1] for _, r in trozo])
        escribe_tabla(doc, filas, False, alterna=True)
        for n, (ident, _, _) in enumerate(ps, 1):
            if ident in ERRATAS:
                p = doc.add_paragraph(style="Aviso")
                p.add_run("Ojo con la %d: " % n).bold = True
                p.add_run(re.sub(r"</?b>", "", ERRATAS[ident]))
                _recuadra(p)


MUESTRARIO = [
    ("Portada rótulo", "OPOSICIONES RTVE · CONVOCATORIAS 1/2022 Y 3/2022"),
    ("Title", "Temario general"),
    ("Subtitle", "El subtítulo de la portada, a dos renglones si hace falta"),
    ("Portada dato", "Redacción vigente a 21 de diciembre de 2022"),
    ("Encabezado TOAC", "TOAC – Temarios de Oposiciones"),
    ("Pie", "Página 3 de 254        Oposiciones RTVE – Temario General"),
    ("Rótulo", "TEMA 1 DEL TEMARIO GENERAL"),
    ("Heading 1", "Título del tema"),
    ("Heading 2", "1.1 Un epígrafe"),
    ("Heading 3", "1.1.1 Un subepígrafe"),
    ("Heading 4", "Un cuarto nivel, que sale poco"),
    ("Normal", "El cuerpo del tema. Lleva negritas para el dato que se pregunta y "
               "cursivas para los títulos de norma. Es el estilo que decide cómo se lee "
               "el volumen entero, así que es el primero que conviene fijar."),
    ("Aviso", "Los párrafos de la caja «Cómo usar este volumen»."),
    ("Cita de la norma", "«El texto literal del BOE, que va sangrado para que se "
                         "distinga del cuerpo a simple vista.»"),
    ("List Bullet", "Un punto de lista de primer nivel"),
    ("List Bullet 2", "Un punto de lista anidado"),
    ("List Number", "Un punto de lista numerada"),
    ("Pregunta", "1. El enunciado de una pregunta real de examen."),
    ("Opción", "a) Una de las cuatro opciones."),
    ("Fuente de pregunta", "Nota al pie de una pregunta, cuando hace falta"),
    ("Separador", ""),
]


def muestrario(doc):
    """Una hoja con todos los estilos y su nombre delante."""
    salto(doc)
    doc.add_paragraph("Para dar el formato, no forma parte del volumen", style="Rótulo")
    doc.add_paragraph("Muestrario de estilos", style="Heading 1")
    p = doc.add_paragraph(style="Normal")
    p.add_run("Cada renglón va escrito con el estilo cuyo nombre lleva delante. ")
    p.add_run("Cambiando el estilo en Word cambian todas sus apariciones en el volumen")
    p.runs[-1].bold = True
    p.add_run(", y el fichero devuelto dice qué se ha elegido para cada uno. En el "
              "documento no hay formato aplicado a mano en ningún sitio: si algo se ve "
              "distinto es porque tiene un estilo distinto. Esta hoja se puede borrar.")
    for nombre, ejemplo in MUESTRARIO:
        doc.add_paragraph("· %s" % nombre, style="Fuente de pregunta")
        if ejemplo:
            doc.add_paragraph(ejemplo, style=nombre)
        else:
            _raya(doc.add_paragraph(style=nombre))
    doc.add_paragraph("Tabla y Tabla cabecera", style="Fuente de pregunta")
    escribe_tabla(doc, [[md.parse(c)[1] for c in fila] for fila in
                        (("**Columna**", "**Otra columna**"),
                         ("El texto de las celdas lleva el estilo «Tabla»",
                          "y la primera fila, «Tabla cabecera»"))], True)
    doc.add_paragraph(style="Normal")
    p = doc.add_paragraph(style="Pregunta")
    p.add_run("1. ", style="Número de pregunta")
    p.add_run("El número de la pregunta lleva el estilo de carácter «Número de "
              "pregunta»; el texto con acento monoespaciado, ")
    p.add_run("«Código»", style="Código")
    p.add_run(".")


def main():
    args = [a for a in sys.argv[1:] if not a.startswith("--")]
    salida = args[0] if args else "libro-general.docx"
    quiere = None
    for a in sys.argv[1:]:
        if a.startswith("--temas"):
            valor = a.split("=", 1)[1] if "=" in a else sys.argv[sys.argv.index(a) + 1]
            quiere = {int(x) for x in re.findall(r"\d+", valor)}

    doc = prepara(Document())
    elegidos = [(i, base, banco) for i, (base, banco) in enumerate(TEMAS, 1)
                if quiere is None or i in quiere]

    entradas, hechos, total = [], [], 0
    for numero, base, _ in elegidos:
        crudo = sin_marcas(lee("temas/general/%s.md" % base))
        titulo = re.search(r"(?m)^# (.+)$", crudo).group(1)
        cuerpo = re.sub(r"## Índice\n.*?(?=\n## )", "", crudo, flags=re.S)
        entradas.append((numero, titulo, re.findall(r"(?m)^## (.+)$", cuerpo)))
    for _, _, banco in elegidos:
        if banco:
            total += len(preguntas(banco))

    portada(doc, total, len(elegidos))
    cabecera_y_pie(doc)
    aviso(doc)
    indice(doc)
    for numero, base, banco in elegidos:
        ps = parte_tema(doc, numero, base, banco)
        if ps:
            hechos.append((numero, banco, ps))
    respuestas(doc, hechos)
    if "--muestrario" in sys.argv:
        muestrario(doc)

    ruta = os.path.join(RAIZ, salida)
    doc.save(ruta)
    print("· %s · %d KB · %d tema(s) · %d preguntas"
          % (salida, os.path.getsize(ruta) // 1024, len(elegidos), total))


if __name__ == "__main__":
    main()
